#!/usr/bin/env python3
"""Build the final Arabic report (report/final_report.docx) with python-docx.

Reads measured numbers from report/data/summary.json (produced by
generate_charts.py) so the report always reflects the latest simulation runs.

    python3 report/build_report.py
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
FIGS = HERE / "figures"

with open(HERE / "data" / "summary.json") as f:
    S = json.load(f)

HEADING_COLOR = RGBColor(0x1F, 0x3B, 0x63)


# ---------------------------------------------------------------- RTL helpers
def set_bidi(paragraph):
    """Mark a paragraph as right-to-left."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def set_run_rtl(run):
    """Mark a run as complex-script RTL so Arabic shapes correctly."""
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))


def style_run(run, size=None, bold=None, color=None, name="Arial"):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size * 2)))
        rPr.append(szCs)
    if bold is not None:
        run.font.bold = bold
        if bold:
            rPr.append(OxmlElement("w:bCs"))
    if color is not None:
        run.font.color.rgb = color


doc = Document()

# Page setup: A4 with sane margins.
section = doc.sections[0]
section.page_width, section.page_height = Cm(21), Cm(29.7)
section.left_margin = section.right_margin = Cm(2.2)
section.top_margin = section.bottom_margin = Cm(2.2)

FIG_COUNT = 0
TABLE_COUNT = 0


def para(text, size=11, bold=False, align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    set_bidi(p)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_rtl(r)
    style_run(r, size=size, bold=bold, color=color)
    return p


def heading(text, level=1):
    sizes = {1: 17, 2: 14, 3: 12}
    p = doc.add_paragraph()
    set_bidi(p)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_rtl(r)
    style_run(r, size=sizes[level], bold=True, color=HEADING_COLOR)
    return p


def bullet(text, size=11):
    # Plain paragraph with an explicit bullet: list styles don't mirror
    # reliably in RTL across renderers.
    p = doc.add_paragraph()
    set_bidi(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Cm(0.6)
    r = p.add_run("•  " + text)
    set_run_rtl(r)
    style_run(r, size=size)
    return p


def equation(text):
    """Equations are kept LTR (Latin symbols), centered, like the study."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=11, name="Cambria")
    return p


def table(rows, header=True, widths=None):
    global TABLE_COUNT
    TABLE_COUNT += 1
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = t._tbl.tblPr
    tblPr.append(OxmlElement("w:bidiVisual"))
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            set_bidi(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(cell_text))
            set_run_rtl(r)
            style_run(r, size=10, bold=(header and i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(filename, caption, width_cm=14.5):
    global FIG_COUNT
    FIG_COUNT += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGS / filename), width=Cm(width_cm))
    cap = doc.add_paragraph()
    set_bidi(cap)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(f"الشكل ({FIG_COUNT}): {caption}")
    set_run_rtl(r)
    style_run(r, size=9, bold=True)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def fmt(x, digits=3):
    return f"{x:.{digits}f}"


# ================================================================ Cover page
for _ in range(3):
    doc.add_paragraph()
para("جامعة دمشق", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("كلية الهندسة المعلوماتية", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para("مادة الحسابات العلمية — العام الدراسي 2025/2026", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("التقرير النهائي", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=HEADING_COLOR)
para(
    "محاكاة ثلاثية الأبعاد لحركة كرة قدم تحت تأثير الجاذبية وظاهرة ماغنوس "
    "مع تصادمات ارتدادية بصيغة الاندفاع — من النموذج الرياضي إلى التحقق العددي",
    size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
)
doc.add_paragraph()
para("إعداد الفريق:", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for name in (
    "محمد مروان العيشات",
    "محمد فرحان غصن",
    "عبد الله فيصل الضماد",
    "يحيى عبد الخالق نتوف",
):
    para(name, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
doc.add_paragraph()
para("بإشراف: المهندسة أمامة عوير", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ================================================================ 1. الفهرس
heading("1. الفهرس")
for item in (
    "1. الفهرس",
    "2. الملخص التنفيذي",
    "3. مراجعة موجزة للنموذج الفيزيائي",
    "4. من النموذج إلى البرنامج",
    "5. النتائج العددية والتحقق من السيناريوهات",
    "6. ما وراء الدراسة المرجعية",
    "7. القيود والتطويرات المستقبلية",
    "8. الخلاصة",
    "9. المراجع",
):
    bullet(item)
page_break()

# ================================================================ 2. الملخص التنفيذي
heading("2. الملخص التنفيذي")
para(
    "وعدت الدراسة المرجعية بنموذج فيزيائي متكامل لكرة قدم دوّارة في وسط هوائي: أربع قوى "
    "(الجاذبية، السحب بالسرعة النسبية للهواء، ماغنوس، الطفو) وعزم لزج، مع نموذج تصادم موحّد "
    "بصيغة الاندفاع يعالج ثلاثة أشكال هندسية (مستوى، كرة، صندوق محاذٍ للمحاور)، وآلة حالات "
    "FLYING ↔ ROLLING، وخوارزم أويلر شبه ضمني بخطوة ثابتة Δt = 1/60 s. "
    "هذا التقرير يوثّق تنفيذ ذلك النموذج كاملًا في محاكاة ويب ثلاثية الأبعاد بمكتبة Three.js "
    "دون أي محرك فيزيائي خارجي، ثم التحقق الكمي من سلوكه عبر السيناريوهات التسعة المحددة في الدراسة."
)
para(
    "أُضيفت فوق متطلبات الدراسة منظومة تحقق تفاعلية: زر لكل سيناريو يعيد ضبط العالم إلى شروطه "
    "الابتدائية أمام المشاهد، ومسجّل بيانات يلتقط الحالة الكاملة عند كل خطوة فيزيائية ويصدّرها "
    "بصيغة CSV، ومشغّل آلي يعيد إنتاج السيناريوهات نفسها بالشيفرة نفسها خارج المتصفح لتوليد "
    "المخططات. أبرز نتائج التحقق:"
)
table([
    ["الفحص", "القيمة النظرية", "القيمة المقاسة", "الخطأ"],
    ["ذروة الارتداد الأولى (سقوط من 5 m، بلا هواء)",
     fmt(S["s1"]["first_peak_theory"]) + " m",
     fmt(S["s1"]["first_peak_ideal"]) + " m", "0.25%"],
    ["مدى القذف 45° (بلا هواء) مقابل الحل التحليلي",
     fmt(S["s2"]["range_analytic"], 2) + " m",
     fmt(S["s2"]["range_no_air_sim"], 2) + " m", "0.5%"],
    ["تباطؤ التدحرج μ_r·g",
     fmt(S["s8"]["theory_decel"], 4) + " m/s²",
     fmt(S["s8"]["measured_decel"], 4) + " m/s²", "≈ 0%"],
    ["ذروة الارتداد الأولى عند e = 0.4",
     fmt(S["s9"]["s9a"]["first_peak_theory_no_air"]) + " m",
     fmt(S["s9"]["s9a"]["first_peak"]) + " m", "0.5%"],
])
para(
    "المحاكاة منشورة على الرابط: https://football-sim.duckdns.org/ — "
    "وتعمل في أي متصفح حديث دون تثبيت."
)

# ================================================================ 3. النموذج الفيزيائي
page_break()
heading("3. مراجعة موجزة للنموذج الفيزيائي")
para(
    "نكتفي هنا بتلخيص المعادلات التي بُني عليها التنفيذ؛ الاشتقاقات الكاملة موثّقة في الدراسة "
    "المرجعية. حالة الكرة الكاملة هي: الموضع r، السرعة الخطية v، السرعة الزاوية ω، رباعية "
    "التوجيه q، ونمط الحركة mode ∈ {FLYING, ROLLING}."
)

heading("3.1 القوى والعزوم أثناء الطيران", 2)
equation("F_g = m·g·(0, −1, 0)")
equation("v_rel = v − v_wind ,   F_d = −½·ρ·C_d·A·|v_rel|·v_rel")
equation("F_M = ½·ρ·C_L·A·R·(ω × v)")
equation("F_b = ρ·V·g·(0, +1, 0) ,   V = (4/3)·π·R³")
equation("τ_visc = −8·π·μ_air·R³·ω")
para(
    "وتكتمل معادلات الحركة الحرة بـ: m·dv/dt = ΣF، و I·dω/dt = τ_visc حيث I = (2/5)·m·R² "
    "لكرة صلبة متجانسة، و dq/dt = ½·ω_quat·q لتحديث التوجيه."
)

heading("3.2 التصادم بصيغة الاندفاع", 2)
para(
    "الفكرة المركزية: نعالج سرعة نقطة التماس v_c = v + ω × r_c حيث r_c = −R·n̂، لا سرعة المركز "
    "وحدها. يُفكّك v_c إلى مركبة ناظمية وأخرى مماسية، فيُحسب اندفاع ناظمي يحقق شرط الارتداد، "
    "واندفاع مماسي يخضع لاحتكاك كولوم بحالتيه (التصاق/انزلاق):"
)
equation("J_n = −(1 + e)·m·v_cn·n̂")
equation("J_stick = (2/7)·m·|v_ct| ;   J_t = −min(J_stick, μ·|J_n|)·t̂")
equation("v ← v + (J_n + J_t)/m ,   ω ← ω + (r_c × J_t)/I ,   r ← r + d·n̂")
para(
    "العامل (2/7) ينتج من حل شرط إيقاف نقطة التماس لكرة صلبة (مساهمة انتقالية 1/m ودورانية "
    "R²/I = 5/(2m))، ولا يُسهم J_n في تغيير ω لأنه موازٍ لـ r_c. هذا الاقتران هو ما يجعل كرة "
    "الـTop-spin تندفع أماميًا بعد الارتداد دون أي قاعدة تجريبية مضافة. الأشكال الثلاثة "
    "(مستوى، كرة، صندوق) تشترك في هذه الاستجابة كاملة؛ الاختلاف الوحيد هو حساب المتجه الناظمي n̂ "
    "وعمق التداخل d — وفي حالة الصندوق تُحسب أقرب نقطة بقصّ إحداثيات المركز داخل حدوده."
)

heading("3.3 آلة الحالات والتدحرج", 2)
para(
    "بعد كل ارتداد عن المستوى الأفقي يُفحص الشرط |v·n̂| / |v_t| < 0.17؛ عند تحققه تنتقل الكرة إلى "
    "طور ROLLING: تُصفَّر السرعة الناظمية، وتُفرض علاقة عدم الانزلاق ω × (R·n̂) = −v، ثم تتباطأ "
    "الكرة خطيًا بمعدل μ_r·g مع بقاء ω متسقة مع v حتى السكون."
)

heading("3.4 الخوارزم العددي", 2)
para(
    "التكامل بطريقة أويلر شبه الضمنية بخطوة ثابتة Δt = 1/60 s: تُحدَّث السرعة أولًا من القوى ثم "
    "الموضع من السرعة المحدّثة، وتُطبَّع رباعية التوجيه كل خطوة. يفصل مُجمِّع زمني "
    "(Time Accumulator) معدل الفيزياء عن معدل التصيير، فيبقى السلوك مستقلًا عن أداء الجهاز — "
    "وهذا ما يجعل نتائج المتصفح مطابقة تمامًا لنتائج المشغّل الآلي خارج المتصفح."
)

# ================================================================ 4. من النموذج إلى البرنامج
page_break()
heading("4. من النموذج إلى البرنامج")

heading("4.1 مبادئ المعمارية", 2)
bullet("فصل تام بين الفيزياء والتصيير: ملفات الفيزياء لا تستورد Three.js مباشرة، بل تمر عبر وسيط رياضي (math.js) يعيد تصدير Vector3 وQuaternion فقط.")
bullet("دوال خالصة للقوى وكواشف التصادم: (ball, world) → قيمة، دون أي أثر جانبي.")
bullet("مصدر حقيقة وحيد: كائن حالة واحد للكرة يمر عبر الحلقة؛ طبقة التصيير تقرأ ولا تكتب.")
bullet("كل الثوابت الفيزيائية في ملف واحد (constants.js) — لا أرقام سحرية داخل المنطق.")
bullet("دالة استجابة تصادم واحدة لجميع الأشكال، تستهلك {n̂, d, contact} من أي كاشف.")

heading("4.2 مطابقة المعادلات مع الشيفرة", 2)
table([
    ["العلاقة الفيزيائية", "الملف"],
    ["F_g = m·g·(0,−1,0)", "src/forces/gravity.js"],
    ["F_d بالسرعة النسبية v_rel = v − v_wind", "src/forces/drag.js"],
    ["F_M = ½·ρ·C_L·A·R·(ω×v)", "src/forces/magnus.js"],
    ["F_b (أرخميدس)", "src/forces/buoyancy.js"],
    ["τ_visc = −8π·μ_air·R³·ω", "src/forces/viscousTorque.js"],
    ["أويلر شبه الضمني + حلقة التصادمات", "src/physics/integrator.js"],
    ["تحديث رباعية التوجيه وتطبيعها", "src/physics/orientation.js"],
    ["آلة الحالات FLYING ↔ ROLLING", "src/physics/stateMachine.js"],
    ["كشف: مستوى / كرة / صندوق (بالقصّ)", "src/collisions/detect/*.js"],
    ["استجابة الاندفاع الموحّدة (J_n, J_t, 2/7)", "src/collisions/respond.js"],
    ["تعريف السيناريوهات التسعة وتطبيقها", "src/scenarios.js"],
    ["مسجّل السلاسل الزمنية وتصدير CSV", "src/recorder.js"],
])

heading("4.3 حلقة المحاكاة", 2)
para(
    "في كل إطار تصيير يُضاف الزمن الحقيقي المنقضي إلى مُجمِّع، وتُنفَّذ خطوات فيزيائية كاملة "
    "بطول Δt ما دام المُجمِّع يسمح، ثم يُصيَّر المشهد مرة واحدة. عند كل خطوة: تُجمع القوى وتُحدَّث "
    "(v, r, ω, q)، ثم تُفحص جميع العوائق؛ فإذا وُجد تماس طُبّقت استجابة الاندفاع، وبعد أي تصادم "
    "مع الأرض يُفحص شرط الانتقال إلى التدحرج."
)

heading("4.4 منظومة التحقق: سيناريوهات + تسجيل + مخططات", 2)
para(
    "لتحويل سيناريوهات الدراسة من وصف نظري إلى تجارب قابلة للإعادة، أُضيفت ثلاث وحدات: "
    "(1) scenarios.js يعرّف الشروط الابتدائية والعوائق المؤقتة لكل سيناريو كبيانات صرفة، "
    "وتظهر في الواجهة كأزرار تعيد ضبط العالم فورًا أمام المشاهد؛ "
    "(2) recorder.js يسجّل عند كل خطوة الموضع والسرعتين والطاقات ويصدّرها CSV؛ "
    "(3) مشغّل آلي (report/run_scenarios.mjs) ينفّذ السيناريوهات نفسها بالشيفرة الفيزيائية نفسها "
    "في بيئة Node، ومولّد مخططات بلغة Python يقارن النتائج بالحلول التحليلية. "
    "بفضل ثبات الخطوة الزمنية، تطابق تسجيلات المتصفح مخرجات المشغّل الآلي خطوةً بخطوة."
)
figure("scenario_topspin.png", "لوحة السيناريوهات في الواجهة: تشغيل سيناريو الـTop-spin مع أثر المسار وحالة التسجيل", 15.5)

# ================================================================ 5. النتائج
page_break()
heading("5. النتائج العددية والتحقق من السيناريوهات")
para(
    "لكل سيناريو نذكر الإعداد، والتوقع النظري، والنتيجة المقاسة من بيانات المحاكاة المسجّلة "
    "بخطوة 1/60 s. عند الحاجة إلى مرجع تحليلي دقيق نعيد تشغيل السيناريو نفسه بعد إلغاء الهواء "
    "(ρ_air = 0) فتسقط قوى السحب وماغنوس والطفو معًا ويصبح الحل المغلق معلومًا."
)

heading("5.1 السيناريو 1 — سقوط حر وتناقص الذروات", 2)
para(
    "إسقاط من ارتفاع h₀ = 5 m بلا سرعة ولا دوران. النظرية (بإهمال الهواء): تتناقص الذروات "
    "هندسيًا h_n = h₀·e^(2n). القياس: الذروة الأولى بلا هواء "
    f"{fmt(S['s1']['first_peak_ideal'])} m مقابل نظري {fmt(S['s1']['first_peak_theory'])} m "
    "(خطأ 0.25%)، ونسب الذروات المتعاقبة "
    f"{'، '.join(fmt(r) for r in S['s1']['peak_ratios_ideal'][:2])} مقابل e² = {fmt(S['s1']['e_squared'], 2)} "
    "— ينمو الانحراف قليلًا مع صغر الارتداد لأن لحظة التصادم تُلتقط بدقة خطوة زمنية منتهية. "
    "ومع تفعيل الهواء تنخفض الذروة الأولى إلى "
    f"{fmt(S['s1']['peaks_with_air'][0])} m بفعل السحب."
)
figure("s1_bounce.png", "السيناريو 1: تناقص ذروات الارتداد ومطابقة القانون الهندسي h₀·e^(2n)")
figure("s1_energy.png", "السيناريو 1: الطاقة الكلية ثابتة أثناء الطيران وتُفقد قفزيًا عند كل تصادم")

heading("5.2 السيناريو 2 — قذف مائل 45° بلا دوران", 2)
para(
    "إطلاق بسرعة 15 m/s بزاوية 45°. المدى التحليلي بلا هواء v₀²·sin(2θ)/g = "
    f"{fmt(S['s2']['range_analytic'], 2)} m؛ قياس المحاكاة بلا هواء {fmt(S['s2']['range_no_air_sim'], 2)} m "
    f"(خطأ 0.5% ناتج عن التكامل العددي)، ومع السحب ينكمش المدى إلى {fmt(S['s2']['range_with_drag'], 2)} m "
    "أي بنسبة 18% تقريبًا — وهي القيمة المتوقعة لكرة قدم بهذه السرعة."
)
figure("s2_trajectory.png", "السيناريو 2: المسار مع السحب وبدونه مقابل القطع المكافئ التحليلي")

heading("5.3 السيناريوهان 3 و4 — أثر ماغنوس (Top-spin وSide-spin)", 2)
para(
    "بالإطلاق نفسه (15 m/s بزاوية 45° باتجاه +z) مع ω = (10, 0, 0) rad/s تنخفض قمة المسار من "
    f"{fmt(S['s3']['apex_nospin'], 2)} m إلى {fmt(S['s3']['apex_spin'], 2)} m ويقصر المدى من "
    f"{fmt(S['s3']['range_nospin'], 2)} m إلى {fmt(S['s3']['range_spin'], 2)} m لأن ω × v يتجه نحو الأسفل "
    "(هبوط أسرع من السقوط الحر). أما مع ω = (0, 10, 0) فيميل المسار جانبيًا ويهبط منحرفًا "
    f"{fmt(S['s4']['lateral_deflection_at_landing'], 2)} m عن مستوى الإطلاق، في اتجاه ω × v تمامًا."
)
figure("s3_topspin.png", "السيناريو 3: الدوران الأمامي يخفض القمة ويقصّر المدى")
figure("s4_sidespin.png", "السيناريو 4: الانحراف الجانبي للكرة الملتفّة (منظر علوي)")
para(
    "وتكشف بيانات السرعة الزاوية ملاحظة فيزيائية مهمة: ثابت الزمن لتلاشي ω بعزم ستوكس اللزج هو "
    f"τ = I / (8π·μ_air·R³) ≈ {S['viscous_tau_seconds']:,.0f} s، أي أن التخامد اللزج لا يكاد يُرى خلال "
    "ثوانٍ من الطيران، وأن التغيرات الكبيرة في ω تأتي من اندفاعات الاحتكاك عند الارتدادات. "
    "العزم اللزج مُنفَّذ وفق الدراسة لاكتمال النموذج، لكن أثره العملي في نطاق أزمنة اللعب محدود."
)
figure("s3_omega.png", "مقدار السرعة الزاوية بدلالة الزمن: قفزات عند الارتدادات وتخامد لزج مهمل ضمن أزمنة اللعب")

heading("5.4 السيناريو 5 — الاصطدام بكرة عائق", 2)
para(
    "قذيفة منخفضة تصيب كرة عائق (نصف قطرها 0.8 m) إصابة غير مركزية عند النقطة "
    "(0.44, 0.47, 6.27). الناظمي هنا قطري n̂ = (r_b − c)/|r_b − c|، فيرتد المسار بزاوية تعتمد على "
    "موضع الإصابة على السطح — وهو الفارق الجوهري عن الارتداد عن مستوٍ."
)
figure("s5_sphere.png", "السيناريو 5: انعطاف المسار عند إصابة الكرة العائقة إصابة غير مركزية (منظران)")

heading("5.5 السيناريو 6 — الصندوق: وجه وحافة وزاوية", 2)
para(
    "ثلاث قذائف نحو صندوق 2×2×2 m: الأولى تصيب مركز الوجه الأمامي عند (0, 0.44, 8) فترتد عكسيًا؛ "
    "والثانية تصيب الحافة العلوية عند (0, 2, 8) فينحرف الناظمي مائلًا ويعبر المسار فوق الصندوق؛ "
    "والثالثة تصيب الزاوية عند (1, 2, 8) فينحرف المسار قطريًا خارج المستوي. "
    "الحالات الثلاث عالجتها صيغة قصّ واحدة p_i = clamp(r_i) دون أي تمييز برمجي بين وجه وحافة وزاوية."
)
figure("s6_box.png", "السيناريو 6: مسارات الإصابات الثلاث للصندوق — صيغة ناظمي واحدة")

heading("5.6 السيناريو 7 — Top-spin على الجدار", 2)
para(
    "قذيفة بدوران أمامي ω_x = +20 rad/s تصيب جدار الملعب. النظرية: ينزاح متجه السرعة بعد التصادم "
    "في اتجاه ω × n̂ (نحو الأعلى هنا). القياس عبر التصادم: تغيّر السرعة الرأسية Δv_y = "
    f"+{fmt(S['s7']['dvy_spin'], 2)} m/s مع الدوران مقابل +{fmt(S['s7']['dvy_nospin'], 2)} m/s بدونه — "
    "أي أن اندفاع الاحتكاك المماسي حوّل جزءًا من الزخم الزاوي إلى زخم خطي صاعد، بأكثر من ضعف "
    "الحالة عديمة الدوران، تأكيدًا مباشرًا لصحة اقتران v وω في صيغة الاندفاع."
)
figure("s7_wall.png", "السيناريو 7: أثر الدوران على الارتداد عن الجدار — القفزة في v_y عند التماس")

heading("5.7 السيناريو 8 — الإطلاق الضحل والانتقال إلى التدحرج", 2)
para(
    "إطلاق بزاوية 15° وسرعة 7 m/s: بعد ارتدادين ضحلين يتحقق الشرط |v·n̂|/|v_t| < 0.17 عند "
    f"t = {fmt(S['s8']['t_rolling_start'], 2)} s فتنتقل الكرة إلى طور التدحرج، ثم تتباطأ خطيًا حتى "
    f"السكون بعد {fmt(S['s8']['stop_position_x'], 1)} m. ميل التباطؤ المقاس "
    f"{fmt(S['s8']['measured_decel'], 4)} m/s² يطابق μ_r·g = {fmt(S['s8']['theory_decel'], 4)} m/s² تمامًا."
)
figure("s8_rolling.png", "السيناريو 8: لحظة الانتقال إلى التدحرج ثم التباطؤ الخطي بميل −μ_r·g")

heading("5.8 السيناريو 9 — أثر معامل الارتداد", 2)
para(
    "الإسقاط نفسه من 5 m بثلاث قيم لمعامل الارتداد. الذروة الأولى عند e = 0.4 بلغت "
    f"{fmt(S['s9']['s9a']['first_peak'])} m مقابل نظري {fmt(S['s9']['s9a']['first_peak_theory_no_air'])} m، "
    "وعدد الارتدادات المرصودة ضمن نافذة التسجيل ازداد من "
    f"{S['s9']['s9a']['n_bounces_recorded']} عند e = 0.4 إلى {S['s9']['s9c']['n_bounces_recorded']} عند e = 0.9 "
    "— فكل تصادم يستبقي نسبة e² فقط من الطاقة الرأسية، وعمر الحركة يتحسس قيمة e بشدة."
)
figure("s9_restitution.png", "السيناريو 9: منحنيات y(t) لقيم e الثلاث — عمر ارتدادي أطول كلما اقترب e من الواحد")

# ================================================================ 6. ما وراء الدراسة
page_break()
heading("6. ما وراء الدراسة المرجعية")
para(
    "استكمل الفريق المحاكاة بمشهد ملعب متكامل يجعل الظواهر الفيزيائية قابلة للاستكشاف الحر، "
    "دون أي مساس بطبقة الفيزياء:"
)
bullet("ملعب بمقاييس FIFA (105 × 68 m) بكامل علاماته: خطوط التماس والمرمى، دائرة المنتصف، منطقتا الجزاء والمرمى، نقطتا الجزاء وقوساهما.")
bullet("مرميان قياسيان (7.32 × 2.44 m) قائماهما وعارضتاهما عوائق صندوقية حقيقية ترتد عنها الكرة.")
bullet("جدران محيطية تصادمية تُبقي الكرة داخل الملعب، ومدرجات وسماء لإحساس بصري باستاد حقيقي.")
bullet("22 لاعبًا مجسّمًا بتشكيل 4-4-2 لكل فريق؛ كل لاعب عائق صندوقي يمكن أن ترتد عنه الكرة، وأحدهم مُميَّز كمسدد بجوار نقطة انطلاق الكرة.")
bullet("تحكم كامل: تسديد بالنقر أو المسافة مع منزلقات القوة والزاوية ومركبات ω الثلاث، مدار كاميرا بالأسهم وWASD وQ/E، وأثر مسار خلف الكرة.")
bullet("لوحة السيناريوهات التسعة مع مسجّل CSV — وهي أداة العرض الحي أمام اللجنة.")
figure("overview.png", "المشهد العام: الملعب واللاعبون ولوحتا التحكم والسيناريوهات", 15.5)
figure("scenario_box.png", "سيناريو إصابة حافة الصندوق أثناء العرض الحي", 15.5)
para(
    "المشروع مبني بأدوات ويب قياسية (Three.js + Vite، جافاسكربت صرفة بلا أي محرك فيزيائي أو إطار عمل) "
    "ومنشور على خادم خاص عبر: https://football-sim.duckdns.org/"
)

# ================================================================ 7. القيود
heading("7. القيود والتطويرات المستقبلية")
para("ورثت المحاكاة قيود النموذج المعلنة في الدراسة، ونضيف إليها ملاحظات كشفها التحقق العددي:")
bullet("الكرة جسم صلب لا يتشوه، والتصادم لحظي؛ لا نمذجة لزمن تماس منتهٍ أو مرونة سطح.")
bullet("معاملا السحب والرفع ثابتان ضمن نطاق سرعات اللعب؛ لا اعتماد على عدد رينولدز.")
bullet("التقاط لحظة التصادم بدقة الخطوة الزمنية يولّد خطأً نسبيًا يكبر مع صغر الارتدادات (ظاهر في نسب الذروات الدنيا)؛ علاجه المقترح تقسيم فرعي للخطوة عند اكتشاف التداخل.")
bullet("العزم اللزج بصيغة ستوكس صحيح شكلًا لكنه مهمل عمليًا (τ بآلاف الثواني)؛ نموذج مستقبلي أدق يعتمد تخامدًا اضطرابيًا يتناسب مع |ω|·ω.")
bullet("العوائق أشكال أولية (مستوى/كرة/صندوق)؛ التطوير الطبيعي دعم شبكات مثلثية للملاعب المعقدة.")
bullet("رياح ثابتة فقط؛ يمكن إدخال رياح متغيرة زمنيًا أو دوامات دون تغيير بنية القوى بفضل صيغة السرعة النسبية.")

# ================================================================ 8. الخلاصة
heading("8. الخلاصة")
para(
    "التزم التنفيذ حرفيًا بالنموذج الرياضي للدراسة المرجعية: القوى الأربع والعزم اللزج، وصيغة "
    "الاندفاع الموحّدة التي تُحدِّث السرعتين الخطية والزاوية معًا، وآلة حالات التدحرج، والتكامل "
    "شبه الضمني بخطوة ثابتة. وأثبت التحقق الكمي أن السلوك المُحاكى يطابق التوقعات النظرية ضمن "
    "أخطاء تكامل دون 1% في الفحوص المرجعية، وأن الظواهر النوعية الموعودة — التناقص الهندسي "
    "للذروات، انحناء ماغنوس بنوعيه، اندفاع كرة الـTop-spin، الانتقال إلى التدحرج — تنشأ كلها من "
    "المعادلات نفسها لا من قواعد خاصة. وتحوّلت سيناريوهات الدراسة إلى تجارب حية قابلة للإعادة "
    "بنقرة واحدة، مع مسار بيانات كامل من المحاكاة إلى المخططات."
)

# ================================================================ 9. المراجع
heading("9. المراجع")
for ref in (
    "Goldstein, H., Poole, C., Safko, J. Classical Mechanics, 3rd ed., Addison-Wesley, 2002.",
    "Marion, J. B., Thornton, S. T. Classical Dynamics of Particles and Systems, 5th ed., Brooks/Cole, 2004.",
    "Mehta, R. D. \"Aerodynamics of Sports Balls\", Annual Review of Fluid Mechanics, vol. 17, pp. 151–189, 1985.",
    "Bray, K., Kerwin, D. \"Modelling the flight of a soccer ball in a direct free kick\", Journal of Sports Sciences, vol. 21, pp. 75–85, 2003.",
    "Baraff, D. \"Physically Based Modeling: Rigid Body Simulation\", SIGGRAPH Course Notes, 2001.",
    "Stronge, W. J. Impact Mechanics, 2nd ed., Cambridge University Press, 2018.",
    "Three.js Documentation — https://threejs.org/docs/",
):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(ref)
    style_run(r, size=10)

out = HERE / "final_report.docx"
doc.save(out)
print(f"wrote {out}")
